import sys, os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
from config.paths import *

"""
generate_word_with_full_legends.py

Generates a Word document with two tables and full, comprehensive legends
(including definitions, statistical methods, and abbreviations).
"""

import pandas as pd
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ============================================================================
# Paths
# ============================================================================
DATA_DIR = os.path.join(DIABETES_DIR, "data", "methylation")
TABLES_DIR = os.path.join(DIABETES_DIR, "results", "tables")
OUTPUT_DOCX = os.path.join(DIABETES_DIR, "results", "Tables_with_Full_Legends.docx")

# ============================================================================
# Load data
# ============================================================================
print("Loading data...")

# Expression table
expr_df = pd.read_csv(os.path.join(TABLES_DIR, "Table1_Expression_Summary.csv"))

# Methylation table
methyl_df = pd.read_csv(os.path.join(DATA_DIR, "methylation_final_stats.csv"))
methyl_df['mean_disease'] = methyl_df['mean_disease'].fillna(methyl_df['mean_t2d'])
methyl_df = methyl_df[['condition', 'motif', 'mean_control', 'mean_disease', 'fold_change', 'p_value']]
methyl_df.columns = ['Condition', 'Motif', 'Mean_Control', 'Mean_Disease', 'Fold_Change', 'p_value']
methyl_df['Mean_Control'] = methyl_df['Mean_Control'].round(3)
methyl_df['Mean_Disease'] = methyl_df['Mean_Disease'].round(3)
methyl_df['Fold_Change'] = methyl_df['Fold_Change'].round(3)

# Add significance stars
def stars(p):
    if pd.isna(p):
        return ""
    if p < 0.001:
        return "***"
    elif p < 0.01:
        return "**"
    elif p < 0.05:
        return "*"
    else:
        return ""

expr_df['Significance'] = expr_df['p_value'].apply(stars)
methyl_df['Significance'] = methyl_df['p_value'].apply(stars)

# ============================================================================
# Create Word Document
# ============================================================================
print("Creating Word document with full legends...")
doc = Document()

# Title
doc.add_heading('Supplementary Tables for Manuscript', 0)
doc.add_paragraph('This document contains the final statistical tables with comprehensive legends for the analysis of conserved non-CG motifs in pancreatic beta cells across normal, T1D, and T2D conditions.')

# ============================================================================
# Table 1
# ============================================================================
doc.add_heading('Table 1: Expression Summary in Pancreatic Beta Cells', level=1)

# Table
table = doc.add_table(rows=1, cols=len(expr_df.columns))
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
hdr_cells = table.rows[0].cells
for i, col in enumerate(expr_df.columns):
    hdr_cells[i].text = col
    for p in hdr_cells[i].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(10)

# Data rows
for idx, row in expr_df.iterrows():
    cells = table.add_row().cells
    for i, col in enumerate(expr_df.columns):
        val = row[col]
        if isinstance(val, float):
            if col == 'p_value':
                cells[i].text = f"{val:.4f}"
            else:
                cells[i].text = f"{val:.3f}"
        elif isinstance(val, bool):
            cells[i].text = "Yes" if val else "No"
        else:
            cells[i].text = str(val)
        for p in cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(9)

# Full Legend for Table 1
legend1 = (
    "Legend: Target genes (n = 10) were defined as protein‑coding genes located within 50 kb of the 11 conserved "
    "non‑CG motifs on chromosome X. Control genes (n = 16) were selected from chromosome X at a distance of more than "
    "500 kb from any conserved motif, with matched gene size and expression coverage. Scaled expression values (range 0–1) "
    "were obtained from the CZ CELLxGENE platform (March 2025 release) using publicly available single‑cell RNA‑sequencing "
    "datasets. Group means were compared using independent‑samples t‑tests (Welch’s correction for unequal variance). "
    "Fold change is calculated as Target_Mean / Control_Mean. Significance thresholds: * p < 0.05; ** p < 0.01; *** p < 0.001. "
    "p‑values were adjusted for multiple testing using Bonferroni correction. Abbreviations: T1D, type 1 diabetes; "
    "T2D, type 2 diabetes; SD, standard deviation; N, number of genes."
)
doc.add_paragraph(legend1, style='Normal')
doc.add_paragraph()

# ============================================================================
# Table 2
# ============================================================================
doc.add_heading('Table 2: DNA Methylation Summary at Conserved Non-CG Motifs (T1D and T2D)', level=1)

# Table
table = doc.add_table(rows=1, cols=len(methyl_df.columns))
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
hdr_cells = table.rows[0].cells
for i, col in enumerate(methyl_df.columns):
    hdr_cells[i].text = col
    for p in hdr_cells[i].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(10)

# Data rows
for idx, row in methyl_df.iterrows():
    cells = table.add_row().cells
    for i, col in enumerate(methyl_df.columns):
        val = row[col]
        if isinstance(val, float):
            if col == 'p_value':
                cells[i].text = f"{val:.4f}"
            else:
                cells[i].text = f"{val:.3f}"
        else:
            cells[i].text = str(val)
        for p in cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(9)

# Full Legend for Table 2
legend2 = (
    "Legend: Beta‑values (ranging from 0 to 1) were extracted from Illumina 450K or 27K methylation arrays for probes "
    "overlapping the conserved non‑CG motifs. For T1D, data were obtained from GSE124809 (IFNα‑treated human islets, 450K); "
    "for T2D, data were obtained from GSE21232 (pancreatic islets, 27K). Mean beta values are shown for control and disease "
    "groups. Fold change is calculated as Mean_Disease / Mean_Control. Statistical significance was assessed using Welch’s "
    "t‑test (unequal variance). Significance thresholds: * p < 0.05; ** p < 0.01; *** p < 0.001. For T2D, four motifs "
    "(AAGA, GAGG, TATG, ATCT) showed a significant reduction in methylation (fold change ≈ 0.53, p = 0.0034), indicated by "
    "double asterisks (**). No significant methylation changes were observed in T1D (all p ≈ 1.0). Abbreviations: T1D, "
    "type 1 diabetes; T2D, type 2 diabetes; IFNα, interferon‑alpha; N, number of samples."
)
doc.add_paragraph(legend2, style='Normal')
doc.add_paragraph()

# ============================================================================
# Save
# ============================================================================
doc.save(OUTPUT_DOCX)
print(f"SUCCESS: Word document with full legends saved to: {OUTPUT_DOCX}")
print(f"  - Table 1: {len(expr_df)} rows")
print(f"  - Table 2: {len(methyl_df)} rows")